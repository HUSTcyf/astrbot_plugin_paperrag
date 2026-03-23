"""
Paper RAG Plugin - 核心RAG引擎模块
实现本地论文库的检索增强生成功能
"""

import asyncio
import json
import os
import sys
import time
import warnings
import inspect
from pathlib import Path
from typing import List, Dict, Any, Optional, cast, Union, TYPE_CHECKING
from dataclasses import dataclass, field
from datetime import datetime

from astrbot.api import logger
from astrbot.core.provider.provider import EmbeddingProvider  # type: ignore

# Ollama类型导入（TYPE_CHECKING避免循环导入）
if TYPE_CHECKING:
    from .ollama_embedding import OllamaEmbeddingProvider
else:
    OllamaEmbeddingProvider = None  # type: ignore

# 抑制底层库的 gRPC/absl 警告（这些是正常的，不影响功能）
os.environ['GRPC_VERBOSITY'] = 'ERROR'  # 只显示 gRPC 错误
os.environ['GLOG_minloglevel'] = '2'     # 抑制 Google 日志

try:
    from pymilvus import connections, utility, Collection, CollectionSchema, FieldSchema
    from pymilvus.client.types import DataType as MilvusDataType  # 使用别名避免冲突
    import fitz  # PyMuPDF
except ImportError as e:
    raise ImportError(f"请先安装必要依赖: pip install pymilvus PyMuPDF")

# 导入语义分块模块
try:
    from .semantic_chunker import PDFParserAdvanced, SemanticChunker, Chunk
    SEMANTIC_CHUNKER_AVAILABLE = True
except ImportError as e:
    SEMANTIC_CHUNKER_AVAILABLE = False
    # 定义类型占位符，避免类型检查错误
    PDFParserAdvanced = None  # type: ignore
    SemanticChunker = None  # type: ignore
    Chunk = None  # type: ignore
    logger.warning(f"⚠️ 语义分块模块不可用，将使用基础分块方式: {e}")

# 导入重排序模块
try:
    from .reranker import create_reranker, AdaptiveReranker
    RERANKER_AVAILABLE: bool = True
except ImportError as e:
    RERANKER_AVAILABLE = False
    # 类型占位符，避免Pylance报告未绑定变量
    AdaptiveReranker = None  # type: ignore
    create_reranker = None  # type: ignore
    logger.warning(f"⚠️ 重排序模块不可用: {e}")

# 可选的多格式支持
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("⚠️ python-docx 未安装，Word文档支持将被禁用")


@dataclass
class RAGConfig:
    """RAG配置类"""
    # Provider配置
    embedding_provider_id: str = ""
    llm_provider_id: str = ""

    # Embedding模式配置
    embedding_mode: str = "api"  # "api" 或 "ollama"
    ollama_config: dict = field(default_factory=dict)  # Ollama配置

    # Milvus配置 - 支持 Lite 模式和远程服务器
    milvus_lite_path: str = "./data/milvus_papers.db"  # Lite 模式路径
    address: str = ""  # 远程 Milvus 服务器地址
    db_name: str = "default"  # 数据库名称
    authentication: Optional[dict] = None  # 认证信息 {token, user, password}
    collection_name: str = "paper_embeddings"

    # 检索配置
    embed_dim: int = 768  # Gemini默认768维
    top_k: int = 5
    similarity_cutoff: float = 0.3

    # 论文目录
    papers_dir: str = "./papers"

    # 语义分块配置
    chunk_size: int = 512  # 目标块大小
    chunk_overlap: int = 0  # 重叠大小（默认0，避免无限循环bug）
    min_chunk_size: int = 100  # 最小块大小
    use_semantic_chunking: bool = True  # 是否使用语义分块

    # 多模态配置
    enable_multimodal: bool = True  # 是否启用多模态提取

    # 重排序配置
    enable_reranking: bool = False  # 是否启用重排序（提升精度，但增加延迟）
    reranking_model: str = "BAAI/bge-reranker-v2-m3"  # 重排序模型
    reranking_device: str = "auto"  # 重排序设备（auto/mps/cuda/cpu）
    reranking_adaptive: bool = True  # 自适应重排序（智能决定是否重排序）
    reranking_threshold: float = 0.0  # 重排序分数阈值
    reranking_batch_size: int = 32  # 重排序批处理大小

    def __post_init__(self):
        """初始化后处理"""
        if self.authentication is None:
            self.authentication = {}
        if self.ollama_config is None:
            self.ollama_config = {}

    def validate(self) -> tuple[bool, str]:
        """验证配置"""
        if self.embed_dim % 64 != 0:
            return False, "嵌入维度必须是64的倍数"

        # 检查 Milvus 配置
        if not self.milvus_lite_path and not self.address:
            # 两个都为空，使用默认 Lite 路径
            self.milvus_lite_path = "./data/milvus_papers.db"

        return True, ""

    def get_connection_mode(self) -> str:
        """获取连接模式：'lite' 或 'remote'"""
        # milvus_lite_path 优先级更高
        if self.milvus_lite_path and self.milvus_lite_path.strip():
            return 'lite'
        elif self.address and self.address.strip():
            return 'remote'
        else:
            # 默认使用 Lite 模式
            return 'lite'

    def get_connection_uri(self) -> str:
        """获取连接 URI"""
        mode = self.get_connection_mode()
        if mode == 'lite':
            return self.milvus_lite_path
        else:
            return self.address


class EmbeddingProviderWrapper:
    """AstrBot Embedding Provider 包装类 - 支持 OpenAI、Gemini 等API Provider"""

    def __init__(self, provider: Any):
        if not provider:
            raise ValueError("Embedding provider 不能为 None")
        self.provider = provider

    async def embed(self, texts: str | List[str]) -> List[List[float]]:
        """批量获取文本嵌入（优先使用批量API，自动分批以符合API限制）"""
        try:
            if isinstance(texts, str):
                texts = [texts]

            # Gemini API限制：单次批量请求最多100个
            BATCH_SIZE_LIMIT = 100

            # 辅助函数：智能调用（同步/异步自适应）
            async def call_method(method, *args):
                if inspect.iscoroutinefunction(method):
                    return await method(*args)
                else:
                    return method(*args)

            # 辅助函数：执行单批次embedding
            async def embed_batch(batch: List[str]) -> List[List[float]]:
                """处理单个批次的embedding"""
                # 优先级：批量方法 > 单个方法
                # 1. get_embeddings (Gemini/OpenAI批量方法)
                if hasattr(self.provider, "get_embeddings") and callable(self.provider.get_embeddings):
                    return await call_method(self.provider.get_embeddings, batch)

                # 2. embed_texts (其他Provider的批量方法)
                elif hasattr(self.provider, "embed_texts") and callable(self.provider.embed_texts):
                    return await call_method(self.provider.embed_texts, batch)

                # 3. embed (通用embed方法)
                elif hasattr(self.provider, "embed") and callable(self.provider.embed):
                    return await call_method(self.provider.embed, batch)

                # 4. get_embedding (单个方法，逐个调用 - 兜底方案)
                elif hasattr(self.provider, "get_embedding") and callable(self.provider.get_embedding):
                    return [await call_method(self.provider.get_embedding, text) for text in batch]
                else:
                    available_methods = [
                        attr for attr in ['get_embeddings', 'embed_texts', 'embed', 'get_embedding']
                        if hasattr(self.provider, attr) and callable(getattr(self.provider, attr))
                    ]
                    raise Exception(f"Provider 没有支持的 embedding 方法。可用方法: {available_methods}")

            # 分批处理以符合API限制
            if len(texts) > BATCH_SIZE_LIMIT:
                logger.info(f"📦 文本数量 ({len(texts)}) 超过批量限制，将分批处理（每批 {BATCH_SIZE_LIMIT} 个）")
                all_embeddings = []
                for i in range(0, len(texts), BATCH_SIZE_LIMIT):
                    batch = texts[i:i + BATCH_SIZE_LIMIT]
                    batch_num = i // BATCH_SIZE_LIMIT + 1
                    total_batches = (len(texts) + BATCH_SIZE_LIMIT - 1) // BATCH_SIZE_LIMIT
                    logger.debug(f"🔄 处理批次 {batch_num}/{total_batches} ({len(batch)} 个文本)")
                    batch_embeddings = await embed_batch(batch)
                    all_embeddings.extend(batch_embeddings)
                return all_embeddings
            else:
                logger.debug(f"🚀 使用批量embed方法处理 {len(texts)} 个文本（1次API调用）")
                embeddings = await embed_batch(texts)

            if not embeddings:
                raise Exception("Embedding provider 返回空结果")

            return embeddings

        except Exception as e:
            logger.error(f"Embedding 调用失败: {e}")
            raise Exception(f"获取 embedding 失败: {e}")

    async def get_text_embedding(self, text: str) -> List[float]:
        result = await self.embed([text])
        return result[0] if result and len(result) > 0 else []

    async def get_query_embedding(self, query: str) -> List[float]:
        """获取查询嵌入"""
        return await self.get_text_embedding(query)


class MilvusStore:
    """Milvus 向量存储封装（使用旧式 API，避免阻塞）"""

    def __init__(self, uri: str, collection_name: str, dim: int, db_name: str = "default", authentication: Optional[dict] = None):
        self.uri = uri
        self.collection_name = collection_name
        self.dim = dim
        self.db_name = db_name
        self.authentication = authentication or {}
        self._alias = f"paper_rag_{collection_name}"  # 唯一的连接别名
        self._is_connected = False
        self._collection_initialized = False

    def _ensure_connected(self):
        """确保已连接到 Milvus（延迟连接）"""
        if self._is_connected:
            return

        try:
            connect_params = {}

            if self.uri.endswith(".db"):
                import os
                os.makedirs(os.path.dirname(self.uri) or ".", exist_ok=True)
                connect_params["uri"] = self.uri
            else:
                connect_params["uri"] = self.uri

            if self.authentication.get("token"):
                connect_params["token"] = self.authentication["token"]
            elif self.authentication.get("user"):
                connect_params["user"] = self.authentication["user"]
                if self.authentication.get("password"):
                    connect_params["password"] = self.authentication["password"]

            if self.db_name and self.db_name != "default":
                connect_params["db_name"] = self.db_name

            connections.connect(alias=self._alias, **connect_params)
            self._is_connected = True

        except Exception as e:
            logger.error(f"Milvus 连接失败: {e}")
            raise

    async def _ensure_collection(self):
        """确保集合已初始化（延迟初始化）"""
        if self._collection_initialized:
            return

        self._ensure_connected()

        try:
            from pymilvus import utility
            has_collection = utility.has_collection(self.collection_name, using=self._alias)

            if not has_collection:
                fields = [
                    FieldSchema(name="id", dtype=MilvusDataType.INT64, is_primary=True, auto_id=True),  # type: ignore[arg-type]
                    FieldSchema(name="vector", dtype=MilvusDataType.FLOAT_VECTOR, dim=self.dim),  # type: ignore[arg-type]
                    FieldSchema(name="text", dtype=MilvusDataType.VARCHAR, max_length=65535),  # type: ignore[arg-type]
                    FieldSchema(name="metadata", dtype=MilvusDataType.JSON)  # type: ignore[arg-type]
                ]

                schema = CollectionSchema(fields=fields, description="paper embeddings collection")
                collection = Collection(name=self.collection_name, schema=schema, using=self._alias)

                is_lite_mode = self.uri.endswith(".db")

                if is_lite_mode:
                    index_params = {"index_type": "AUTOINDEX", "metric_type": "COSINE"}
                else:
                    index_params = {"index_type": "HNSW", "metric_type": "COSINE", "params": {"M": 8, "efConstruction": 64}}

                index_result = collection.create_index(field_name="vector", index_params=index_params)

                if asyncio.iscoroutine(index_result):
                    await index_result
                elif index_result is not None and hasattr(index_result, 'done'):
                    index_result.result()  # type: ignore[arg-type]

                load_result = collection.load()
                if load_result is not None and hasattr(load_result, 'done'):
                    load_result.result()  # type: ignore[arg-type]

            self._collection_initialized = True

        except Exception as e:
            logger.error(f"集合初始化失败: {e}")
            raise

    @property
    def client(self):
        """兼容性属性，返回别名"""
        self._ensure_connected()
        return self._alias

    async def add_documents(self, documents: List[Dict[str, Any]]) -> int:
        """批量添加文档"""
        if not documents:
            return 0

        await self._ensure_collection()

        try:
            from pymilvus import Collection

            collection = Collection(self.collection_name, using=self._alias)

            for i, doc in enumerate(documents):
                if asyncio.iscoroutine(doc["embedding"]):
                    raise RuntimeError(f"documents[{i}]['embedding'] 是 coroutine！上游代码忘记 await")
                if not isinstance(doc["embedding"], (list, tuple)):
                    raise TypeError(f"documents[{i}]['embedding'] 类型错误：{type(doc['embedding'])}")

            data = [
                {
                    "vector": doc["embedding"],
                    "text": doc["text"],
                    "metadata": doc.get("metadata", {})
                }
                for doc in documents
            ]

            insert_result = collection.insert(data)
            collection.flush()

            return len(documents)

        except Exception as e:
            logger.error(f"添加文档失败：{e}")
            return 0

    async def search(self, query_vector: List[float], top_k: int = 5,
                     filter_expr: str = '') -> List[Dict[str, Any]]:
        """向量搜索"""
        # 确保集合已初始化
        await self._ensure_collection()

        try:
            from pymilvus import Collection

            # 获取集合
            collection = Collection(self.collection_name, using=self._alias)

            # 执行搜索
            results = collection.search(
                data=[query_vector],
                anns_field="vector",
                param={"metric_type": "COSINE", "params": {"ef": 64}},
                limit=top_k,
                output_fields=["text", "metadata"]
            )

            # search() 可能返回协程或 SearchFuture
            if asyncio.iscoroutine(results):
                results = await results
            elif results is not None and hasattr(results, 'done'):
                results = results.result()  # type: ignore[arg-type]

            # 确保结果不为空
            if results is None:
                return []

            # 格式化结果
            formatted_results = []
            for result in results[0]:  # type: ignore[index]  # 第一个查询的结果
                formatted_results.append({
                    "text": result.entity.get("text", ""),
                    "metadata": result.entity.get("metadata", {}),
                    "score": result.distance
                })

            return formatted_results
        except Exception as e:
            logger.error(f"❌ 搜索失败: {e}")
            return []

    async def list_papers(self) -> List[Dict[str, Any]]:
        """列出所有论文"""
        await self._ensure_collection()

        try:
            from pymilvus import Collection

            collection = Collection(self.collection_name, using=self._alias)
            results = collection.query(expr="", output_fields=["metadata"], limit=16384)

            if asyncio.iscoroutine(results):
                results = await results
            elif results is not None and hasattr(results, 'done'):
                results = results.result()  # type: ignore[arg-type]

            papers = {}
            for result in results:
                metadata = result.entity.get("metadata", {})
                filename = metadata.get("file_name", "unknown")

                if filename not in papers:
                    papers[filename] = {
                        "file_name": filename,
                        "chunk_count": 0,
                        "added_time": metadata.get("added_time", "")
                    }
                papers[filename]["chunk_count"] += 1

            return list(papers.values())
        except Exception as e:
            logger.error(f"列出论文失败: {e}")
            return []

    async def clear(self):
        """清空集合"""
        try:
            from pymilvus import utility

            self._ensure_connected()

            if utility.has_collection(self.collection_name, using=self._alias):
                # drop_collection可能返回None、协程或Future，需要特殊处理
                result = utility.drop_collection(self.collection_name, using=self._alias)

                # 如果返回的不是None，尝试await
                if result is not None:
                    # 使用辅助函数避免类型收窄问题
                    async def _await_if_possible(obj):
                        """如果对象支持await，则await它"""
                        if asyncio.iscoroutine(obj):
                            await obj
                        elif hasattr(obj, '__await__'):
                            await obj

                    await _await_if_possible(result)

                self._collection_initialized = False
                logger.info(f"✅ 集合 {self.collection_name} 已清空")

        except Exception as e:
            logger.error(f"清空集合失败: {e}")
            raise


class PaperRAGEngine:
    """论文RAG引擎（完全异步初始化模式）"""

    def __init__(self, config: RAGConfig, context):
        self.config = config
        self.context = context

        # 完全延迟初始化：所有组件都延迟到第一次使用时
        self._store: MilvusStore | None = None
        self._store_initialized = False

        self._embedding_provider = None
        # Embedding wrapper可以是API wrapper或Ollama provider
        self._embedding_wrapper: Union[EmbeddingProviderWrapper, "OllamaEmbeddingProvider", None] = None
        self._embedding_initialized = False

        self._llm_provider = None
        self._llm_initialized = False

        # 重排序器
        self._reranker = None
        self._reranker_initialized = False

        self._initialization_failed = False
        self._initialization_error = None

    def _ensure_store_initialized(self):
        """确保 Milvus 存储已初始化（延迟初始化）"""
        if self._store_initialized:
            return

        try:
            uri = self.config.get_connection_uri()
            if uri.endswith(".db"):
                Path(uri).parent.mkdir(parents=True, exist_ok=True)

            self._store = MilvusStore(
                uri=uri,
                collection_name=self.config.collection_name,
                dim=self.config.embed_dim,
                db_name=self.config.db_name,
                authentication=self.config.authentication
            )
            self._store_initialized = True
        except Exception as e:
            logger.error(f"Milvus 存储初始化失败: {e}")
            self._initialization_failed = True
            self._initialization_error = str(e)
            raise

    def _ensure_embedding_initialized(self):
        """确保 Embedding Provider 已初始化（延迟初始化）"""
        if self._embedding_initialized:
            return

        try:
            # Ollama模式：使用本地Ollama服务
            if self.config.embedding_mode == "ollama":
                from .ollama_embedding import create_ollama_provider

                ollama_cfg = self.config.ollama_config or {}
                base_url = ollama_cfg.get("base_url", "http://localhost:11434")
                model = ollama_cfg.get("model", "bge-m3")
                timeout = ollama_cfg.get("timeout", 120.0)
                batch_size = ollama_cfg.get("batch_size", 10)
                retry_attempts = ollama_cfg.get("retry_attempts", 3)

                logger.info(f"🦙 使用Ollama本地Embedding（模型: {model}, 地址: {base_url}）")
                self._embedding_wrapper = create_ollama_provider(
                    base_url=base_url,
                    model=model,
                    timeout=timeout,
                    batch_size=batch_size,
                    retry_attempts=retry_attempts
                )

                # 自动更新embed_dim为Ollama模型的维度
                if model == "bge-m3":
                    self.config.embed_dim = 1024
                elif model == "nomic-embed-text":
                    self.config.embed_dim = 768
                elif model == "mxbai-embed-large":
                    self.config.embed_dim = 1024
                else:
                    # 其他模型默认1024维
                    self.config.embed_dim = 1024

                self._embedding_initialized = True
                return

            # API模式：使用配置的API Provider
            if not self.config.embedding_provider_id:
                raise Exception("API模式下必须配置embedding_provider_id")

            provider_manager = getattr(self.context, "provider_manager", None)
            if provider_manager is None:
                raise Exception("无法访问 context.provider_manager")

            inst_map = getattr(provider_manager, "inst_map", None)
            if not isinstance(inst_map, dict):
                raise Exception("inst_map 不是 dict")

            provider = inst_map.get(self.config.embedding_provider_id)
            if provider is None:
                raise Exception(f"Provider '{self.config.embedding_provider_id}' 不存在")

            self._embedding_provider = provider
            self._embedding_wrapper = EmbeddingProviderWrapper(provider)
            self._embedding_initialized = True

        except Exception as e:
            logger.error(f"Embedding Provider 初始化失败: {e}")
            self._initialization_failed = True
            self._initialization_error = str(e)
            raise

    def _ensure_llm_initialized(self):
        """确保 LLM Provider 已初始化（延迟初始化）"""
        if self._llm_initialized or not self.config.llm_provider_id:
            return

        try:
            provider_manager = getattr(self.context, "provider_manager", None)
            inst_map = getattr(provider_manager, "inst_map", None)
            if isinstance(inst_map, dict):
                provider = inst_map.get(self.config.llm_provider_id)
                if provider:
                    self._llm_provider = provider
                    self._llm_initialized = True
        except Exception as e:
            logger.warning(f"⚠️ LLM Provider 初始化失败: {e}")

    def _ensure_reranker_initialized(self):
        """确保重排序器已初始化（延迟初始化）"""
        if self._reranker_initialized or not self.config.enable_reranking:
            return

        if not RERANKER_AVAILABLE:
            logger.warning("⚠️ 重排序功能不可用，FlagEmbedding未安装")
            self._reranker_initialized = True  # 标记为已初始化避免重复尝试
            return

        # 类型断言：RERANKER_AVAILABLE为True时，create_reranker必定可用
        assert create_reranker is not None  # type: ignore

        try:
            self._reranker = create_reranker(
                model_name=self.config.reranking_model,
                device=self.config.reranking_device,
                batch_size=self.config.reranking_batch_size,
                use_fp16=True,
                adaptive=self.config.reranking_adaptive
            )
            self._reranker_initialized = True
            logger.info(f"✅ 重排序器已初始化: {self.config.reranking_model} (batch_size={self.config.reranking_batch_size})")
        except Exception as e:
            logger.error(f"❌ 重排序器初始化失败: {e}")
            self._reranker_initialized = True  # 标记为已初始化避免重复尝试

    @property
    def store(self) -> MilvusStore:
        """获取 Milvus 存储（延迟初始化）"""
        self._ensure_store_initialized()
        assert self._store is not None
        return self._store

    @property
    def embedding(self) -> Union[EmbeddingProviderWrapper, "OllamaEmbeddingProvider"]:
        """获取 Embedding Provider（延迟初始化，API或Ollama模式）"""
        self._ensure_embedding_initialized()
        assert self._embedding_wrapper is not None
        return self._embedding_wrapper

    @property
    def llm(self):
        """获取 LLM Provider（延迟初始化）"""
        self._ensure_llm_initialized()
        return self._llm_provider

    @property
    def reranker(self):
        """获取重排序器（延迟初始化）"""
        self._ensure_reranker_initialized()
        return self._reranker

    async def ingest_papers(self, file_paths: List[str]):
        """导入论文到知识库（异步生成器，yield进度更新）

        Args:
            file_paths: 文件路径列表

        Yields:
            进度更新字典: {
                "current": 当前文件索引,
                "total": 总文件数,
                "filename": 文件名,
                "status": 状态 (parsing/embedding/storing/done/skipped/error),
                "chunks": 片段数（仅done状态）,
                "error": 错误信息（仅error状态）
            }
        """
        if not file_paths:
            return

        start_time = time.time()
        total_chunks = 0
        skipped_count = 0
        error_count = 0

        for idx, file_path in enumerate(file_paths, 1):
            try:
                filename = str(Path(file_path).name)

                yield {
                    "current": idx,
                    "total": len(file_paths),
                    "filename": filename,
                    "status": "parsing"
                }

                chunks = await self._parse_document(file_path)

                if not chunks:
                    skipped_count += 1
                    yield {
                        "current": idx,
                        "total": len(file_paths),
                        "filename": filename,
                        "status": "skipped"
                    }
                    continue

                yield {
                    "current": idx,
                    "total": len(file_paths),
                    "filename": filename,
                    "status": "embedding"
                }

                embeddings = await self.embedding.embed(chunks)

                documents = []
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    documents.append({
                        "text": chunk,
                        "embedding": embedding,
                        "metadata": {
                            "file_name": filename,
                            "file_path": file_path,
                            "chunk_index": i,
                            "added_time": datetime.now().isoformat()
                        }
                    })

                yield {
                    "current": idx,
                    "total": len(file_paths),
                    "filename": filename,
                    "status": "storing"
                }

                count = await self.store.add_documents(documents)
                total_chunks += count

                yield {
                    "current": idx,
                    "total": len(file_paths),
                    "filename": filename,
                    "status": "done",
                    "chunks": count
                }

            except Exception as e:
                error_count += 1
                logger.error(f"导入失败 {Path(file_path).name}: {e}")
                yield {
                    "current": idx,
                    "total": len(file_paths),
                    "filename": str(Path(file_path).name),
                    "status": "error",
                    "error": str(e)
                }

        # 返回最终统计
        elapsed_time = time.time() - start_time

        yield {
            "status": "complete",
            "files_count": len(file_paths),
            "skipped_count": skipped_count,
            "error_count": error_count,
            "chunks_count": total_chunks,
            "elapsed_time": elapsed_time
        }

    def _detect_file_type(self, file_path: str) -> str:
        """检测文件类型"""
        suffix = Path(file_path).suffix.lower()
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.txt': 'text',
            '.md': 'text',
            '.markdown': 'text',
            '.html': 'html',
            '.htm': 'html',
            '.rtf': 'rich_text',
            '.odt': 'odt',
        }
        return type_map.get(suffix, 'unknown')

    async def _parse_pdf(self, file_path: str) -> List[str]:
        """解析PDF文件（使用语义感知分块）"""
        try:
            filename = str(Path(file_path).name)

            # 优先使用语义分块器
            if SEMANTIC_CHUNKER_AVAILABLE and self.config.use_semantic_chunking:
                parser = PDFParserAdvanced(  # type: ignore
                    enable_multimodal=self.config.enable_multimodal
                )
                chunker = SemanticChunker(  # type: ignore
                    chunk_size=self.config.chunk_size,
                    overlap=self.config.chunk_overlap,
                    min_chunk_size=self.config.min_chunk_size,
                    max_chunk_size=self.config.chunk_size * 2
                )

                chunks_dict = parser.parse_and_chunk(file_path, chunker)
                text_chunks = [c["text"] for c in chunks_dict]

                if len(text_chunks) > 50:
                    logger.warning(f"⚠️ {filename}: 生成了 {len(text_chunks)} 个分块（超过50个），可能会影响检索性能")

                return text_chunks

            else:
                # 回退到基础 PyMuPDF 解析
                doc: fitz.Document = fitz.open(file_path)  # type: ignore
                text_chunks: List[str] = []

                for page in doc:  # type: ignore[arg-type]
                    text: str = page.get_text()  # type: ignore
                    for i in range(0, len(text), self.config.chunk_size):
                        chunk: str = text[i:i+self.config.chunk_size]
                        if len(chunk.strip()) > self.config.min_chunk_size:
                            text_chunks.append(chunk.strip())

                doc.close()

                if len(text_chunks) == 0:
                    logger.warning(f"⚠️ {filename}: 没有提取到任何文本（可能是扫描版PDF）")
                elif len(text_chunks) > 50:
                    logger.warning(f"⚠️ {filename}: 生成了 {len(text_chunks)} 个分块（超过50个），可能会影响检索性能")

                return text_chunks

        except Exception as e:
            logger.error(f"PDF解析失败 {file_path}: {e}")
            return []

    async def _parse_docx(self, file_path: str) -> List[str]:
        """解析Word文档（DOCX）"""
        if not DOCX_SUPPORT:
            logger.error(f"❌ Word文档支持未启用，请安装 python-docx")
            return []

        try:
            from docx import Document as DocxDocument
            doc: DocxDocument = DocxDocument(file_path)  # type: ignore
            text_chunks = []
            full_text = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        full_text.append(row_text)

            combined_text = "\n".join(full_text)

            # 分块处理
            chunk_size = 512
            for i in range(0, len(combined_text), chunk_size):
                chunk = combined_text[i:i+chunk_size]
                if len(chunk.strip()) > 50:
                    text_chunks.append(chunk.strip())

            if len(text_chunks) > 50:
                logger.warning(f"⚠️ [DOCX PARSE] {str(Path(file_path).name)}: 生成了 {len(text_chunks)} 个分块（超过50个），可能会影响检索性能")

            return text_chunks

        except Exception as e:
            logger.error(f"❌ Word文档解析失败 {file_path}: {e}")
            # 返回空列表
            return []

    async def _parse_text(self, file_path: str) -> List[str]:
        """解析纯文本文件（TXT, MD等）"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            text_chunks = []
            chunk_size = 512
            for i in range(0, len(text), chunk_size):
                chunk = text[i:i+chunk_size]
                if len(chunk.strip()) > 50:
                    text_chunks.append(chunk.strip())

            if len(text_chunks) > 50:
                logger.warning(f"⚠️ [TEXT PARSE] {str(Path(file_path).name)}: 生成了 {len(text_chunks)} 个分块（超过50个），可能会影响检索性能")

            return text_chunks

        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    text = f.read()

                text_chunks = []
                chunk_size = 512
                for i in range(0, len(text), chunk_size):
                    chunk = text[i:i+chunk_size]
                    if len(chunk.strip()) > 50:
                        text_chunks.append(chunk.strip())

                if len(text_chunks) > 50:
                    logger.warning(f"⚠️ [TEXT PARSE] {str(Path(file_path).name)}: 生成了 {len(text_chunks)} 个分块（超过50个），可能会影响检索性能")

                return text_chunks
            except Exception as e:
                logger.error(f"❌ 文本文件解析失败 {file_path}: {e}")
                return []
        except Exception as e:
            logger.error(f"❌ 文本文件解析失败 {file_path}: {e}")
            return []

    async def _parse_document(self, file_path: str) -> List[str]:
        """通用文档解析器 - 自动检测文件类型并选择合适的解析方法"""
        file_type = self._detect_file_type(file_path)

        if file_type == 'pdf':
            return await self._parse_pdf(file_path)
        elif file_type == 'docx':
            return await self._parse_docx(file_path)
        elif file_type in ('text', 'html', 'markdown'):
            return await self._parse_text(file_path)
        else:
            logger.error(f"❌ 不支持的文件类型: {file_type}")
            return []

    async def search(self, query: str, mode: str = "rag") -> Dict[str, Any]:
        """搜索论文"""
        # 生成查询嵌入
        query_embedding = await self.embedding.get_query_embedding(query)

        # 检索相关文档（获取更多候选结果供重排序）
        retrieval_k = self.config.top_k * 3 if self.config.enable_reranking else self.config.top_k
        results = await self.store.search(
            query_vector=query_embedding,
            top_k=retrieval_k
        )

        # 重排序（如果启用）
        if self.config.enable_reranking and self.reranker and results:
            try:
                results = await self.reranker.rerank(
                    query=query,
                    results=results,
                    top_k=self.config.top_k
                )
                logger.debug(f"🔄 重排序完成: {len(results)} 个结果")
            except Exception as e:
                logger.warning(f"⚠️ 重排序失败，使用原始顺序: {e}")

        if mode == "retrieve":
            return {
                "type": "retrieve",
                "sources": results
            }

        # RAG模式：生成回答
        if not self.llm:
            return {
                "type": "error",
                "message": "LLM未配置，无法生成回答"
            }

        # 构建上下文
        context = "\n\n".join([
            f"[来源{i+1}] {r['metadata'].get('file_name', 'unknown')}\n{r['text']}"
            for i, r in enumerate(results)
        ])

        # 生成回答
        prompt = f"""基于以下论文内容回答问题：

{context}

问题：{query}

请提供详细的回答，并引用相关文献。"""

        try:
            # 使用 AstrBot LLM Provider
            response = await self.llm._call(prompt)
            answer = response.get_content() if hasattr(response, 'get_content') else str(response)

            return {
                "type": "rag",
                "answer": answer,
                "sources": results
            }
        except Exception as e:
            return {
                "type": "error",
                "message": f"生成回答失败: {e}"
            }

    async def list_papers(self) -> List[Dict[str, Any]]:
        """列出所有论文"""
        return await self.store.list_papers()
